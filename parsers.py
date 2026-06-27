"""
Document Parsers for MultiMind AI Enterprise Knowledge Assistant.

Provides:
- PDF text extraction (PyPDF2)
- DOCX text extraction (python-docx)
- Excel/CSV text extraction (openpyxl)
- Unified document ingestion pipeline
- Metadata extraction (author, creation date, etc.)
"""

import os
import io
import json
import logging
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field, asdict

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Represents a parsed document with text and metadata."""
    filename: str
    file_type: str
    text: str
    pages: int = 1
    metadata: Dict = field(default_factory=dict)
    chunks: List[str] = field(default_factory=list)
    user_id: Optional[str] = None
    category: str = "public"

    def to_dict(self) -> Dict:
        data = asdict(self)
        return data


class PDFParser:
    """Extract text from PDF files."""

    def __init__(self):
        self.available = self._check_availability()

    def _check_availability(self) -> bool:
        try:
            import PyPDF2
            return True
        except ImportError:
            logger.warning("[Parser] PyPDF2 not available. Install with: pip install PyPDF2")
            return False

    def parse(self, file_content: bytes, filename: str = "document.pdf") -> ParsedDocument:
        """Parse PDF from bytes."""
        if not self.available:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
        
        import PyPDF2
        
        pdf_file = io.BytesIO(file_content)
        reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        metadata = {}
        
        # Extract metadata
        if reader.metadata:
            metadata = {
                "title": reader.metadata.get("/Title", ""),
                "author": reader.metadata.get("/Author", ""),
                "subject": reader.metadata.get("/Subject", ""),
                "creator": reader.metadata.get("/Creator", ""),
                "producer": reader.metadata.get("/Producer", ""),
            }
        
        # Extract text from all pages
        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"[Parser] Failed to extract text from page {page_num}: {e}")
        
        full_text = "\n\n".join(text_parts)
        
        return ParsedDocument(
            filename=filename,
            file_type="pdf",
            text=full_text,
            pages=len(reader.pages),
            metadata=metadata,
        )


class DOCXParser:
    """Extract text from DOCX files."""

    def __init__(self):
        self.available = self._check_availability()

    def _check_availability(self) -> bool:
        try:
            import docx
            return True
        except ImportError:
            logger.warning("[Parser] python-docx not available. Install with: pip install python-docx")
            return False

    def parse(self, file_content: bytes, filename: str = "document.docx") -> ParsedDocument:
        """Parse DOCX from bytes."""
        if not self.available:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        from docx import Document
        
        doc = Document(io.BytesIO(file_content))
        
        text_parts = []
        metadata = {}
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        # Extract core properties
        try:
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
        except Exception as e:
            logger.debug(f"[Parser] Could not extract DOCX metadata: {e}")
        
        full_text = "\n\n".join(text_parts)
        
        return ParsedDocument(
            filename=filename,
            file_type="docx",
            text=full_text,
            pages=1,  # DOCX doesn't have explicit pages
            metadata=metadata,
        )


class ExcelParser:
    """Extract text from Excel files (XLSX, XLS)."""

    def __init__(self):
        self.available = self._check_availability()

    def _check_availability(self) -> bool:
        try:
            import openpyxl
            return True
        except ImportError:
            logger.warning("[Parser] openpyxl not available. Install with: pip install openpyxl")
            return False

    def parse(self, file_content: bytes, filename: str = "spreadsheet.xlsx") -> ParsedDocument:
        """Parse Excel from bytes."""
        if not self.available:
            raise ImportError("openpyxl is required for Excel parsing. Install with: pip install openpyxl")
        
        from openpyxl import load_workbook
        
        wb = load_workbook(io.BytesIO(file_content), read_only=True, data_only=True)
        
        text_parts = []
        metadata = {
            "sheet_names": wb.sheetnames,
            "sheet_count": len(wb.sheetnames),
        }
        
        # Extract text from all sheets
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_text = [f"--- Sheet: {sheet_name} ---"]
            
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    sheet_text.append(row_text)
            
            text_parts.extend(sheet_text)
        
        wb.close()
        full_text = "\n".join(text_parts)
        
        return ParsedDocument(
            filename=filename,
            file_type="excel",
            text=full_text,
            pages=len(wb.sheetnames) if hasattr(wb, 'sheetnames') else 1,
            metadata=metadata,
        )


class CSVParser:
    """Extract text from CSV files."""

    def parse(self, file_content: bytes, filename: str = "data.csv") -> ParsedDocument:
        """Parse CSV from bytes."""
        import csv
        
        text_parts = []
        metadata = {}
        
        try:
            # Try to decode as UTF-8
            content_str = file_content.decode('utf-8')
        except UnicodeDecodeError:
            content_str = file_content.decode('latin-1')
        
        reader = csv.reader(io.StringIO(content_str))
        rows = list(reader)
        
        if rows:
            metadata["columns"] = rows[0]
            metadata["row_count"] = len(rows) - 1
        
        for row in rows:
            row_text = " | ".join(row)
            if row_text.strip():
                text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)
        
        return ParsedDocument(
            filename=filename,
            file_type="csv",
            text=full_text,
            pages=1,
            metadata=metadata,
        )


class TextParser:
    """Parse plain text files."""

    def parse(self, file_content: bytes, filename: str = "document.txt") -> ParsedDocument:
        """Parse text from bytes."""
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            text = file_content.decode('latin-1')
        
        return ParsedDocument(
            filename=filename,
            file_type="text",
            text=text,
            pages=1,
            metadata={},
        )


class DocumentParser:
    """Unified document parser that handles multiple file types."""

    def __init__(self):
        self.parsers = {
            "pdf": PDFParser(),
            "docx": DOCXParser(),
            "doc": DOCXParser(),  # python-docx handles both
            "xlsx": ExcelParser(),
            "xls": ExcelParser(),
            "csv": CSVParser(),
            "txt": TextParser(),
            "md": TextParser(),
            "markdown": TextParser(),
        }
        self.chunk_size = 500
        self.chunk_overlap = 50

    def parse_file(self, file_content: bytes, filename: str, user_id: str = None, category: str = "public") -> ParsedDocument:
        """
        Parse a file and return ParsedDocument.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename (used to detect type)
            user_id: User who uploaded the document
            category: Document category for RBAC
        
        Returns:
            ParsedDocument object
        """
        ext = filename.lower().split(".")[-1] if "." in filename else "txt"
        parser = self.parsers.get(ext)
        
        if not parser:
            # Try text parser as fallback
            logger.warning(f"[Parser] No parser for .{ext}, falling back to text parser")
            parser = TextParser()
            ext = "txt"
        
        doc = parser.parse(file_content, filename)
        doc.user_id = user_id
        doc.category = category
        
        # Auto-chunk the text
        doc.chunks = self._chunk_text(doc.text)
        
        logger.info(f"[Parser] Parsed {filename}: {len(doc.text)} chars, {len(doc.chunks)} chunks")
        return doc

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            if chunk.strip():
                chunks.append(chunk.strip())
            start = end - self.chunk_overlap
        
        return chunks

    def get_supported_types(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.parsers.keys())


# Global parser instance
_document_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """Get or create the global document parser."""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
    return _document_parser


def parse_document(file_content: bytes, filename: str, user_id: str = None, category: str = "public") -> ParsedDocument:
    """Convenience function to parse a document."""
    parser = get_document_parser()
    return parser.parse_file(file_content, filename, user_id, category)


def get_supported_document_types() -> List[str]:
    """Convenience function to get supported types."""
    parser = get_document_parser()
    return parser.get_supported_types()
